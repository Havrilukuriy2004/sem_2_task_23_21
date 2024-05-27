import pandas as pd
from docx import Document


file_path = input("Enter a data file path: ")

df_sites = pd.read_excel(file_path, sheet_name='Майданчики')
df_works = pd.read_excel(file_path, sheet_name='Роботи')
df_acts = pd.read_excel(file_path, sheet_name='Акти')
df_items = pd.read_excel(file_path, sheet_name='Пункти')

def create_act_document(act_id, act_no, act_date, act_sum, site_info, work_items):
    doc = Document()
    doc.add_heading(f'Акт № {act_no}', 0)
    doc.add_paragraph(f'Дата: {act_date}')
    doc.add_paragraph(
        f'Даний Акт засвідчує, що Виконавцем на майданчику {site_info["Name"]} були виконані такі роботи:')
    doc.add_paragraph(f'за адресою: {site_info["Address"]}')

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Назва роботи'

    for idx, work_item in enumerate(work_items, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = work_item['Name']

    doc.add_paragraph(f'Сума виконаних робіт складає: {act_sum} грн.')
    doc.add_paragraph('Від Замовника: ____________________')
    doc.add_paragraph('Від Виконавця: ____________________')

    return doc


def generate_acts(df_sites, df_works, df_acts, df_items):
    for _, act in df_acts.iterrows():
        act_id = act['id']
        act_no = act['No']
        act_date = act['Date']
        act_sum = act['Sum']
        site_id = act['S_id']

        site_info = df_sites[df_sites['id'] == site_id].iloc[0]
        work_ids = df_items[df_items['A_id'] == act_id]['W_id']
        work_items = df_works[df_works['id'].isin(work_ids)]

        doc = create_act_document(act_id, act_no, act_date, act_sum, site_info, work_items.to_dict('records'))
        doc.save(f'Act_{act_no}.docx')

generate_acts(df_sites, df_works, df_acts, df_items)

